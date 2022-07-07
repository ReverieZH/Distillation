import json
import numpy as np
import config
from tempfile import TemporaryFile
from flask import Blueprint, _request_ctx_stack, jsonify, request, Response, render_template, send_from_directory, \
    send_file
from sqlalchemy import text, func
from blueprint import RETCODE
from exts import *
from models import ArticleModel
from .decorators import jwt_required
from .utils import JSONEncoder, gen_response_data
from PIL import Image
from docx import Document
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed

bp = Blueprint('article', __name__, url_prefix="/api/article")


@bp.route("/ocr", methods=['POST'])
def generate_by_ocr():
    image_files = request.files.getlist('image[]')
    if not image_files:
        response_data = gen_response_data(RETCODE.PARAMERR, '请上传图片')
        return jsonify(response_data)
    try:
        content = ''
        for image_file in image_files:
            img = Image.open(image_file.stream).convert('RGB')
            img = np.array(img)
            result = ocr.ocr(img)
            for line in result:
                content += line[1][0]
        response_data = gen_response_data(RETCODE.OK, '识别成功', content=content)
    except Exception as e:
        response_data = gen_response_data(RETCODE.EXCEPTION, '识别失败')
    return jsonify(response_data)


@bp.route("/doc", methods=['POST'])
def generate_by_doc():
    doc_file = request.files.get('doc')
    if not doc_file:
        response_data = gen_response_data(RETCODE.PARAMERR, '请上传文档')
        return jsonify(response_data)
    try:
        suffix = doc_file.filename.split(".")[1]
        content = ''
        if suffix == "doc" or suffix == "docx":
            temporary_file = TemporaryFile()
            temporary_file.write(doc_file.stream.read())
            temporary_file.seek(0)
            doc = Document(temporary_file)
            for para in doc.paragraphs:
                content += para.text
        elif suffix == "txt":
            content = doc_file.stream.read().decode(encoding="utf-8")
        elif suffix == "pdf":
            temporary_file = TemporaryFile()
            temporary_file.write(doc_file.stream.read())
            temporary_file.seek(0)
            # 用文件对象来创建一个pdf文档分析器
            praser = PDFParser(temporary_file)
            # 创建一个PDF文档
            doc = PDFDocument()
            # 连接分析器与文档对象
            praser.set_document(doc)
            doc.set_parser(praser)
            # 提供初始化密码，如果没有密码，就创建一个空的字符串
            doc.initialize()
            # 检查文档是否可以转成TXT，如果不可以就忽略
            if not doc.is_extractable:
                raise PDFTextExtractionNotAllowed
            else:
                # 创建PDF资源管理器，来管理共享资源
                rsrcmagr = PDFResourceManager()
                # 创建一个PDF设备对象
                laparams = LAParams()
                # 将资源管理器和设备对象聚合
                device = PDFPageAggregator(rsrcmagr, laparams=laparams)
                # 创建一个PDF解释器对象
                interpreter = PDFPageInterpreter(rsrcmagr, device)
                # 循环遍历列表，每次处理一个page内容
                # doc.get_pages()获取page列表
                for pg in doc.get_pages():
                    interpreter.process_page(pg)
                    # 接收该页面的LTPage对象
                    layout = device.get_result()
                    # 这里的layout是一个LTPage对象 里面存放着page解析出来的各种对象
                    # 一般包括LTTextBox，LTFigure，LTImage，LTTextBoxHorizontal等等一些对像
                    for x in layout:
                        if (isinstance(x, LTTextBox)):  # 网上是判断LTTextBoxHorizontal,而在我写代码的时候，只能判断LTTextBox
                            content += x.get_text()
        response_data = gen_response_data(RETCODE.EXCEPTION, '识别成功', content=content)
    except Exception as e:
        print(e)
        response_data = gen_response_data(RETCODE.EXCEPTION, '识别失败')
    return response_data


@bp.route("/generate", methods=['POST'])
def generate_by_text():
    content = request.json.get('content')
    if not content:
        response_data = gen_response_data(RETCODE.PARAMERR, '请输入文本内容')
        return jsonify(response_data)
    try:
        print(content)
        result = generate.get_title_and_abstract(article=content)
        response_data = gen_response_data(RETCODE.OK, '抽取成功', **result)
    except Exception as e:
        response_data = gen_response_data(RETCODE.EXCEPTION, '抽取失败')
    return jsonify(response_data)


@bp.route("/save", methods=['POST'])
@jwt_required
def save_result():
    user = _request_ctx_stack.top.current_identity
    content = request.json.get('content')
    title = request.json.get('title')
    abstract = request.json.get('abstract')
    if not content or not title or not abstract:
        response_data = gen_response_data(RETCODE.PARAMERR, '请求参数不满足，请输入文章的内容和生成的标题与摘要')
        return jsonify(response_data)
    try:
        history_count = \
            db.session.query(func.count(ArticleModel.id)).filter(ArticleModel.user_id == user['uid']).first()[0]
        filepath = user["username"] + "/" + str(user["uid"]) + "_" + str(history_count + 1) + ".txt"
        article = ArticleModel(title=title, abstract=abstract, filepath=filepath, user_id=user["uid"])
        db.session.add(article)
        response = cos_client.put_object(
            Bucket=config.bucket,
            Body=str(content).encode('utf-8'),
            Key=filepath,
            StorageClass='STANDARD',
            ContentType='text/html; charset=utf-8'
        )
        db.session.commit()
        response_data = gen_response_data(RETCODE.OK, '保存成功', article_id=article.id)
    except Exception as e:
        db.session.rollback()  # 数据库添加失败或云端保存失败时回滚
        print(str(e))
        response_data = gen_response_data(RETCODE.EXCEPTION, '保存失败')
    return jsonify(response_data)


@bp.route("/history", methods=['GET'])
@jwt_required
def get_history():
    user = _request_ctx_stack.top.current_identity
    try:
        articles = ArticleModel.query.filter(ArticleModel.user_id == user['uid'], ArticleModel.status == 1).order_by(
            text('-join_time')).all()
        data_json = json.loads(json.dumps(articles, cls=JSONEncoder))
        response_data = gen_response_data(RETCODE.OK, '查找成功', articles=data_json)
    except Exception as e:
        response_data = gen_response_data(RETCODE.EXCEPTION, '查找失败')
        return jsonify(response_data)
    return jsonify(response_data)


@bp.route("/detail/<id>", methods=['GET'])
@jwt_required
def get_history_detail(id):
    remote_file = TemporaryFile()  # 创建临时文件
    article = ArticleModel.query.filter(ArticleModel.id == id).first()
    if not article:
        response_data = gen_response_data(RETCODE.NODETAIL, '未找到对应详情')
        return jsonify(response_data)
    try:
        # 从云端读取数据
        response = cos_client.get_object(  # 从云端获取对应文本数据
            Bucket=config.bucket,
            Key=article.filepath,
        )
        stream = response['Body'].get_raw_stream().read(1024)
        while stream:
            remote_file.write(stream)
            stream = response['Body'].get_raw_stream().read(1024)
        remote_file.seek(0)
        text = remote_file.read().decode('utf-8')
        # 返回数据
        data_json = json.loads(json.dumps(article, cls=JSONEncoder))
        data_json['text'] = text
        response_data = gen_response_data(RETCODE.OK, '获取成功', article=data_json)
    except Exception as e:
        response_data = gen_response_data(RETCODE.EXCEPTION, '获取失败')
        return jsonify(response_data)
    return response_data


@bp.route("/delete/<id>", methods=['GET'])
@jwt_required
def delete_history(id):
    user = _request_ctx_stack.top.current_identity
    article = ArticleModel.query.filter(ArticleModel.id == id).first()
    if not article:
        response_data = gen_response_data(RETCODE.NODETAIL, '未找到对应历史记录')
        return jsonify(response_data)
    try:
        # 从云端读取数据
        article.status = 0
        db.session.commit()
        articles = ArticleModel.query.filter(ArticleModel.user_id == user['uid'], ArticleModel.status == 1).order_by(
            text('-join_time')).all()
        data_json = json.loads(json.dumps(articles, cls=JSONEncoder))
        response_data = gen_response_data(RETCODE.OK, '删除成功', articles=data_json)
    except Exception as e:
        response_data = gen_response_data(RETCODE.EXCEPTION, '获取失败')
        return jsonify(response_data)
    return jsonify(response_data)


@bp.route('/file', methods=['POST'])
def recive_file():
    content = request.json.get('content')
    title = request.json.get('title')
    abstract = request.json.get('abstract')
    temporary_file = TemporaryFile()
    temporary_file.write(str(title+"\n").encode("utf-8"))
    temporary_file.write(str("\n").encode("utf-8"))
    temporary_file.write(str("abstract\n").encode("utf-8"))
    temporary_file.write(str(abstract+"\n").encode("utf-8"))
    temporary_file.write(str("\n").encode("utf-8"))
    temporary_file.write(str(content).encode("utf-8"))
    temporary_file.seek(0)
    return send_file(temporary_file, download_name="generate.doc", as_attachment=True)

